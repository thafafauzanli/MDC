"""
	Fitur : DistribusiFrekuensi
	Pengarang : Thafa Fauzanli
	Tanggal buat : 17 Maret 2022
"""

import gi
gi.require_version("Gtk", "3.0")
from gi.repository import Gtk
from docx import Document
from docx.shared import Inches
import platform
import webbrowser
import time
import random
import math

class AplDistFre(Gtk.Window):
	def __init__(self, NAMA_APLIKASI, menu, aplvar, jk):
		self.judul = NAMA_APLIKASI + " - Distribusi Frekuensi"
		self.nama = NAMA_APLIKASI
		self.menu = menu
		self.apl = aplvar
		self.lebar = 950
		self.tinggi = 500
		super().__init__(title=self.judul)
		self.set_icon_from_file(jk)
		self.set_default_size(self.lebar, self.tinggi)
		self.set_resizable(False)
		
		# inisialisasi variabel
		self.jumlah_data = 0
		self.nilai_tertinggi = 0
		self.nilai_terendah = 0
		self.j = 0; self.k = 0; self.i = 0
		self.kisi = []
		self.nilai_acak = []
		self.nilai_urut = []
		self.inkelas = []
		self.fre = []
		
		self.pembuatan()

	def pembuatan(self):
		# main kisi (0) =======================================
		self.kisi.append(Gtk.Fixed())
		self.add(self.kisi[0])
		# Inisialisasi(1) ====================================
		self.lInisial = Gtk.Label()
		self.lInisial.set_markup("<big><b>Inisialisasi</b></big>")
		self.fInisial = Gtk.Frame()
		self.fInisial.set_label_widget(self.lInisial)
		self.fInisial.set_shadow_type(Gtk.ShadowType.OUT)
		self.kisi[0].put(self.fInisial, 10, 10)
		self.kisi.append(Gtk.Fixed())
		self.fInisial.add(self.kisi[1])
		self.eJumlah = Gtk.Entry()
		self.eJumlah.set_placeholder_text("Jumlah data")
		self.eJumlah.set_input_purpose(Gtk.InputPurpose.DIGITS)
		self.kisi[1].put(self.eJumlah, 0, 10)
		self.eMin = Gtk.Entry()
		self.eMin.set_placeholder_text("Nilai terendah")
		self.eMin.set_input_purpose(Gtk.InputPurpose.DIGITS)
		self.kisi[1].put(self.eMin, 0, 50)
		self.eMak = Gtk.Entry()
		self.eMak.set_placeholder_text("Nilai tertinggi")
		self.eMak.set_input_purpose(Gtk.InputPurpose.DIGITS)
		self.kisi[1].put(self.eMak, 0, 90)
		self.bAcak = Gtk.Button(label="  Acak  ")
		self.bAcak.connect("clicked", self.data_acak)
		self.kisi[1].put(self.bAcak, 50, 90+40)		
		#  Hasil (2) ==========================================
		self.lHasil = Gtk.Label()
		self.lHasil.set_markup("<big><b>Hasil</b></big>")
		self.fHasil = Gtk.Frame()
		self.fHasil.set_label_widget(self.lHasil)
		self.kisi[0].put(self.fHasil, 200, 10)
		self.kisi.append(Gtk.Fixed())
		self.fHasil.add(self.kisi[2])
		self.lJumlah = Gtk.Label(label="Jumlah data : \n0")
		self.kisi[2].put(self.lJumlah, 0, 10)
		self.lMin = Gtk.Label(label="Nilai terendah : \n0")
		self.kisi[2].put(self.lMin, 120, 10)
		self.lMak = Gtk.Label(label="Nilai tertinggi : \n0")
		self.kisi[2].put(self.lMak, 130*2, 10)
		self.lAcak = Gtk.Label(label="Data acak :")
		self.kisi[2].put(self.lAcak, 0, 50)
		# Nilai Acak
		self.tNila = Gtk.TextBuffer()
		self.tAcak = Gtk.TextView(editable=False, buffer=self.tNila)
		self.tAcak.set_cursor_visible(False)
		self.tAcak.set_size_request(350, 150)
		self.tAcak.set_wrap_mode(Gtk.WrapMode.WORD)
		self.swAcak = Gtk.ScrolledWindow()
		self.swAcak.set_policy(Gtk.PolicyType.AUTOMATIC, Gtk.PolicyType.AUTOMATIC)
		self.swAcak.set_size_request(350, 150)
		self.swAcak.add(self.tAcak)
		self.kisi[2].put(self.swAcak, 0, 70)
		# Nilai Urut
		self.lUrut = Gtk.Label(label="Data terurut :")
		self.kisi[2].put(self.lUrut, 0, 230)
		self.tNilu = Gtk.TextBuffer()
		self.tUrut = Gtk.TextView(editable=False, buffer=self.tNilu)
		self.tUrut.set_cursor_visible(False)
		self.tUrut.set_wrap_mode(Gtk.WrapMode.WORD)
		self.tUrut.set_size_request(350, 150)
		self.swUrut = Gtk.ScrolledWindow()
		self.swUrut.set_policy(Gtk.PolicyType.AUTOMATIC, Gtk.PolicyType.AUTOMATIC)
		self.swUrut.set_size_request(350, 150)
		self.swUrut.add(self.tUrut)
		self.kisi[2].put(self.swUrut, 0, 250)
		# Jarak
		self.lJarak = Gtk.Label(label="Jarak(R) : "+str(self.j))
		self.kisi[2].put(self.lJarak, 380, 10)
		self.tbJarak = Gtk.TextBuffer()
		self.tJarak = Gtk.TextView(editable=False, buffer=self.tbJarak)
		self.tJarak.set_wrap_mode(Gtk.WrapMode.WORD)
		self.tJarak.set_cursor_visible(False)
		self.tJarak.set_size_request(150, 40)
		self.kisi[2].put(self.tJarak, 380, 30)
		# Kelas
		self.lKelas = Gtk.Label(label="Banyak kelas(K) : "+str(self.k))
		self.kisi[2].put(self.lKelas, 550, 10)
		self.tbKelas = Gtk.TextBuffer()
		self.tKelas = Gtk.TextView(editable=False, buffer=self.tbKelas)
		self.tKelas.set_size_request(150, 40)
		self.tKelas.set_cursor_visible(False)
		self.tKelas.set_wrap_mode(Gtk.WrapMode.WORD)
		self.kisi[2].put(self.tKelas, 550, 30)
		self.lInter = Gtk.Label(label="Interval(I) : "+str(self.i))
		self.kisi[2].put(self.lInter, 380, 80)
		self.tbInter = Gtk.TextBuffer()
		self.tInter = Gtk.TextView(editable=False, buffer=self.tbInter)
		self.tInter.set_size_request(150, 40)
		self.tInter.set_wrap_mode(Gtk.WrapMode.WORD)
		self.tInter.set_cursor_visible(False)
		self.kisi[2].put(self.tInter, 380, 100)
		# Tabel Distribusi Frekuensi
		self.fbTabel = Gtk.Frame()
		self.fbTabel.set_size_request(300, 30)
		self.kisi[2].put(self.fbTabel, 380, 160)
		self.tlTabel = Gtk.Label()
		self.tlTabel.set_markup("<big><b>Tabel Distribusi Frekuensi</b></big>")
		self.fbTabel.add(self.tlTabel)
		# Tabel interval kelas
		self.kelas_frame = Gtk.Frame()
		self.kelas_frame.set_size_request(150, 30)
		self.kisi[2].put(self.kelas_frame, 380, 190)
		self.kelas_label = Gtk.Label()
		self.kelas_label.set_markup("<b>Interval Kelas</b>")
		self.kelas_frame.add(self.kelas_label)
		self.kelas_buffer = Gtk.TextBuffer()
		self.kelas_teks = Gtk.TextView(editable=False, buffer=self.kelas_buffer)
		self.kelas_teks.set_justification(Gtk.Justification.CENTER)
		self.kelas_teks.set_size_request(145, 200)
		self.kelas_teks.set_cursor_visible(False)
		# Tabel frekuensi
		self.frekuensi_frame = Gtk.Frame()
		self.frekuensi_frame.set_size_request(150, 30)
		self.kisi[2].put(self.frekuensi_frame, 380+150, 190)
		self.frekuensi_label = Gtk.Label()
		self.frekuensi_label.set_markup("<b>Frekuensi</b>")
		self.frekuensi_frame.add(self.frekuensi_label)
		self.frekuensi_buffer = Gtk.TextBuffer()
		self.frekuensi_teks = Gtk.TextView(editable=False, buffer=self.frekuensi_buffer)
		self.frekuensi_teks.set_cursor_visible(False)
		self.frekuensi_teks.set_justification(Gtk.Justification.CENTER)
		self.frekuensi_teks.set_size_request(145, 200)
		# percobaan scrolling
		self.universal_fixed = Gtk.Layout()
		self.universal_fixed.set_size(150*2, 200)
		self.universal_fixed.set_vexpand(True)
		self.universal_fixed.set_hexpand(True)
		self.universal_scroll = Gtk.ScrolledWindow()
		self.universal_scroll.set_policy(Gtk.PolicyType.AUTOMATIC, Gtk.PolicyType.AUTOMATIC)
		self.universal_scroll.set_size_request(150*2, 200)
		self.universal_scroll.add(self.universal_fixed)
		self.universal_scroll.set_hadjustment(self.universal_fixed.get_hadjustment())
		self.universal_scroll.set_vadjustment(self.universal_fixed.get_vadjustment())
		self.kisi[2].put(self.universal_scroll, 380, 220)
		self.universal_fixed.put(self.kelas_teks,0, 0)
		self.universal_fixed.put(self.frekuensi_teks, 155, 0)
		print(str(self.universal_fixed.get_size()))
		# menu tambahan (3) ==================================================
		self.kisi.append(Gtk.Fixed())
		self.add(self.kisi[3])
		self.lMentam = Gtk.Label()
		self.lMentam.set_markup("<big><b>Alat</b></big>")
		self.mentam = Gtk.Frame()
		self.mentam.set_size_request(160, 250)
		self.mentam.set_label_widget(self.lMentam)
		self.kisi[0].put(self.mentam, 10, 200)
		self.mentam.add(self.kisi[3])
		# tombol bersihkan
		self.tmb_bersihkan = Gtk.Button(label="Bersihkan")
		self.tmb_bersihkan.set_size_request(160, 50)
		self.tmb_bersihkan.connect("clicked", self.bersih)
		self.kisi[3].put(self.tmb_bersihkan, 0, 10)
		# tombol simpan
		self.tmb_simpan = Gtk.Button(label="Simpan docx")
		self.tmb_simpan.connect("clicked", self.simpan)
		self.tmb_simpan.set_size_request(160, 50)
		self.kisi[3].put(self.tmb_simpan, 0, 10+50)
		# tombol keluar
		self.tmb_keluar = Gtk.Button.new_with_label("Kembali")
		self.tmb_keluar.set_size_request(160, 50)
		self.tmb_keluar.connect("clicked", self.kembali)
		self.kisi[3].put(self.tmb_keluar, 0, 10+50*2)
	
	# Fungsi
	def data_acak(self, button):
		if self.eJumlah.get_text() != "" and self.eMin.get_text() != "" and self.eMak.get_text() != "":
			# Membersihkan data
			self.nilai_acak.clear(); self.nilai_urut.clear(); self.inkelas.clear(); self.fre.clear()
			# Mendapatkan data
			self.lJumlah.set_label("Jumlah data : \n" + self.eJumlah.get_text())
			self.lMin.set_label("Nilai terrendah : \n" + self.eMin.get_text())
			self.lMak.set_label("Nilai tertinggi : \n" + self.eMak.get_text())
			jumlah = int(self.eJumlah.get_text())
			mini = int(self.eMin.get_text())
			maksi = int(self.eMak.get_text())
			self.jumlah_data = int(self.eJumlah.get_text())
			self.nilai_terendah = int(self.eMin.get_text())
			self.nilai_tertinggi = int(self.eMak.get_text())
			m = ""; n = ""
			# menciptakan nilai acak
			for i in range(0, jumlah):
				self.nilai_acak.append(random.randint(mini, maksi))
			# mengubah format nilai kebentuk string
			for u in range(0, jumlah):
				if u != jumlah-1:
					m += str(self.nilai_acak[u])+", "
				else:
					m += str(self.nilai_acak[u])
			self.tNila.set_text(m)
			# menciptakan nilai urut
			for i in range(0, jumlah):
				self.nilai_urut.append(self.nilai_acak[i])
			self.nilai_urut.sort()
			for u in range(0, jumlah):
				if u != jumlah-1:
					n += str(self.nilai_urut[u])+", "
				else:
					n += str(self.nilai_urut[u])
			self.tNilu.set_text(n)
			
			# penentuan jarak / range
			self.j = maksi - mini
			self.lJarak.set_label("Jarak(R) : "+str(self.j))
			self.tbJarak.set_text("R = "+str(maksi)+" - "+str(mini)+"\nR = "+str(self.j))
			# penentuan banyak kelas
			self.k = round(1+(3.33*float("{:.2f}".format(math.log10(jumlah)))))
			self.lKelas.set_label("Banyak kelas(K) : "+str(self.k))
			self.tbKelas.set_text(
				"K = 1 + 3.33 log "+str(jumlah)+"\nK = 1 + "+" "+"{:.2f}".format(3.33 * math.log10(jumlah))+
				"\nK = "+str(self.k))
			# penentuan banyak interval
			self.i = round(self.j / self.k)
			self.lInter.set_label("Interval(I) : "+str(self.i))
			self.tbInter.set_text("I = "+str(self.j)+" / "+str(self.k)+"\nI = "+str(self.i))
			# penyusunan interval kelas
			gtw = None
			for i in range(0, self.k):
				if i==0:
					self.inkelas.append(str(mini)+" - "+str(mini+self.i-1))
					gtw = mini+self.i-1
				else:
					gtw = gtw+1
					self.inkelas.append(str(gtw)+" - "+str(gtw+self.i-1))
					gtw = gtw+self.i-1
			gtw1 = ""
			for i in range(0, self.k):
				gtw1 += self.inkelas[i]+"\n"
			self.kelas_buffer.set_text(gtw1)
			gtw = None; gtw1 = ""; gkl = 0 # untuk dibawah
			# penyusunan frekuensi
			for i in range(0, self.k):
				if i==0:
					for u in self.nilai_urut:
						if u in range(mini, mini+self.i):
							gkl += 1
					self.fre.append(gkl); gkl = 0
					gtw = mini+self.i-1
				else:
					gtw += 1
					for u in self.nilai_urut:
						if u in range(gtw, gtw+self.i):
							gkl += 1
					self.fre.append(gkl); gkl = 0
					gtw += self.i-1
			for i in range(0, self.k):
				gtw1 += str(self.fre[i])+"\n"
			self.frekuensi_buffer.set_text(gtw1)
			self.universal_fixed.set_size(150*2, 200+(2*self.k))
							
		else:
			dialog = Gtk.MessageDialog(
				self, 0, Gtk.MessageType.INFO,
				Gtk.ButtonsType.OK, "Inisialisasi nya diisi dong boss..")
			dialog.show()
			dialog.run()
			dialog.destroy()
	
	def bersih(self, tmb):
		# bersihkan inisialisasi
		self.eJumlah.set_text("")
		self.eMin.set_text("")
		self.eMak.set_text("")
		self.jumlah_data = 0
		self.nilai_terendah = 0
		self.nilai_tertinggi = 0
		self.inkelas.clear(); self.fre.clear()
		# bersihkan hasil
		self.j = 0; self.k = 0; self.i = 0
		self.lJumlah.set_label("Jumlah data : \n0")
		self.lMak.set_label("Nilai tertinggi : \n0")
		self.lMin.set_label("Nilai terendah : \n0")
		self.tNila.set_text(""); self.tNilu.set_text("")
		self.lJarak.set_label("Jarak(R) : "+str(self.j))
		self.lKelas.set_label("Banyak kelas(K) : "+str(self.k))
		self.lInter.set_label("Interval(I) : "+str(self.i))
		self.tbJarak.set_text(""); self.tbKelas.set_text(""); self.tbInter.set_text("")
		self.kelas_buffer.set_text(""); self.frekuensi_buffer.set_text("")
		# lainnya
		self.universal_fixed.set_size(150*2, 200)
	
	def simpan(self, tmb):
		if self.inkelas and self.fre:
			# ===== pembuatan docx =====
			p = []
			dok = Document()
			dok.add_heading("Distribusi Frekuensi", 0)
			# paragraf
			p.append(dok.add_paragraph(self.nama)) # 0
			p.append(dok.add_paragraph("Jumlah data = "+str(self.jumlah_data))) # 1
			p.append(dok.add_paragraph("Nilai terkecil = "+str(self.nilai_terendah))) # 2
			p.append(dok.add_paragraph("Nilai terbesar = "+str(self.nilai_tertinggi))) # 3
			p.append(dok.add_paragraph("Data acak = ")) # 4
			for i, d in enumerate(self.nilai_acak):
				if i == len(self.nilai_acak)-1: p[4].add_run(str(d))
				else: p[4].add_run(str(d)+" - ")
			p.append(dok.add_paragraph("Nilai terurut = ")) # 5
			for i, d in enumerate(self.nilai_urut):
				if i == len(self.nilai_urut)-1: p[5].add_run(str(d))
				else: p[5].add_run(str(d)+" - ")
			p.append(dok.add_paragraph("")) # 6
			# tabel
			t = dok.add_table(rows=1, cols=2)
			tc = t.rows[0].cells
			tc[0].text = "Interval Kelas"; tc[1].text = "Frekuensi"
			for kel, fre in zip(self.inkelas, self.fre):
				tc = t.add_row().cells
				tc[0].text = kel
				tc[1].text = str(fre)
			
			# ===== Simpan docx =====
			dok.save("berkas_mdc/ekspor/distribusi_frekuensi"+str(self.jumlah_data)+str(self.nilai_terendah)+str(self.nilai_tertinggi)+".docx")
			dialog = Gtk.MessageDialog(
				self, 0, Gtk.MessageType.INFO,
				Gtk.ButtonsType.OK, "Berhasil menyimpan, cek folder ekspor")
			dialog.show()
			dialog.run()
			dialog.destroy()
		else:
			dialog = Gtk.MessageDialog(
				self, 0, Gtk.MessageType.INFO,
				Gtk.ButtonsType.OK, "Gagal menyimpan, data kosong !")
			dialog.show()
			dialog.run()
			dialog.destroy()
			
	def kembali(self, tmb):
		self.close()
		time.sleep(1)
		self.apl = self.menu()
		self.apl.connect("destroy", Gtk.main_quit)
		self.apl.show_all()
		Gtk.main()
		
