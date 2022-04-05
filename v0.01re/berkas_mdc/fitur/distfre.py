"""
	Fitur : Distribusiself.frekuensi
	Penulis : Thafa Fauzanli
	Tanggal buat : 02 April 2022
"""

import PySimpleGUI as sg
from docx import Document
from docx.shared import Inches
import random
import math
from os.path import exists

class AplDistFre(sg.Window):
	def __init__(self, rm, e):
		# menu
		self.ganti_ruang = e
		self.apl = rm
		col =[]
		nilayout = [[sg.Text("Jumlah data : ", size=(15, 1)),sg.Input("", size=(20,2), focus=False, key="_d")],
		             [sg.Text("nilai terendah :", size=(15,1)), sg.Input("", size=(20,2), focus=False, key="_r")],
		             [sg.Text("nilai tertinggi :", size=(15, 1)), sg.Input("", size=(20,2), focus=False, key="_t")],
		             [sg.Button("Acak"), sg.Button("Bersihkan")]]
		fiturlayout = [[sg.Button("Simpan", size=(10, 3), key="_s"), sg.Button("Kembali", size=(10, 3), key="_k")]]
		hasilayout = [[sg.Text("nilai terendah : 0", size=(15, 1), key="nte"), sg.Text("nilai tertinggi : 0", size=(15, 1), key="nti"), 
					   sg.Text("Jumlah data : 0",size=(15, 1), key="jd")],
					  [sg.Text("Data acak : ")],
					  [sg.Multiline(key="_da", size=(50, 5))],
					  [sg.Text("Data urut : ")],
					  [sg.Multiline(key="_ur", size=(50, 5))],
					  [sg.Frame("Jarak(R)", [[sg.Multiline(key="jarak", size=(20, 5))]]),
					   sg.Frame("Banyak kelas(K)", [[sg.Multiline(key="kelas", size=(20, 5))]]),
					   sg.Frame("Interval(I)", [[sg.Multiline(key="interval", size=(20, 5))]])]]
		self.nilai = []
		kepala = ["Interval Kelas", "frekuensi"]
		tabelayout = [[sg.Table(values=self.nilai, headings=kepala, justification="c", size=(20, 10), key="tabel")]]
		
		# menyiapkan kolom
		col.append(sg.Column([[sg.Frame("Inisialisasi", nilayout)],[sg.Frame("Opsi", fiturlayout)], [sg.Frame("Tabel Distribusi self.frekuensi", tabelayout)]])) # 0
		col.append(sg.Column([[sg.Frame("Hasil", hasilayout)]])) # 1

		# pengaturan jendela
		layout = [[col[0], col[1]]]
		super().__init__("MDC v0.01re - Distribusi frekuensi", layout, icon="berkas_mdc/gambar/ti.png", element_justification="up")
		
		self.ulang()
	
	# event
	def ulang(self):
		nilren, nilting, jumda = 0, 0, 0; pst = ""; pem = None; gem = 0; r = 0; k = 0; i = 0; self.tofre = 0
		self.nila = []; self.nilur = []; self.kelin = []; self.fre = []
		global keruang
		while True:
			event, values = self.read()
			if event in (sg.WIN_CLOSED, "Exit"): # ketika tombol silang ditekan
				break
			if event in ("_k"): # ketika tombol kembali ditekan
				self.close(); self.ganti_ruang(self.apl, None, None)
			if event in ("Acak"): # ketika tombol acak ditekan
				if values["_r"] != "" and values["_d"] != "" and values["_t"] != "":
					# pembersihan
					self.nila.clear(); self.nilur.clear(); self.nilai.clear(); self.kelin.clear(); self.fre.clear(); self.tofre = 0
					# init self.nilai
					nilren = int(values["_r"]); nilting = int(values["_t"]); jumda = int(values["_d"])
					# update
					self["nte"].update(value="nilai terendah : "+values["_r"])
					self["nti"].update(value="nilai tertinggi : "+values["_t"])
					self["jd"].update(value="Jumlah data  : "+values["_d"])
					# pembuatan self.nilai acak
					for i in range(0, jumda):
						self.nila.append(random.randint(nilren, nilting))
					for i in range(0, jumda):
						if i != jumda-1: pst += str(self.nila[i])+", "
						else: pst += str(self.nila[i])
					self["_da"].update(value=pst); pst = ""
					# pengurutan self.nilai dari terkecil - terbesar
					self.nilur = self.nila; self.nilur.sort()
					for i in range(0, jumda):
						if i != jumda-1: pst += str(self.nilur[i])+", "
						else: pst += str(self.nilur[i])
					self["_ur"].update(value=pst); pst = ""
					# menentukan jarak(R)
					j = nilting - nilren
					self["jarak"].update(value="R = "+str(nilting)+" + "+str(nilren)+"\nR = "+str(j))
					# menentukan banyak kelas(K)
					k = round(1+3.33*float("{:.2f}".format(math.log10(jumda))))
					self["kelas"].update(value="K = 1 + 3.33 log "+str(jumda)+"\nK = 1 + 3.33 "+str("{:.2f}".format(math.log10(jumda)))+"\nK = 1 + "+"{:.2f}".format(3.33*math.log10(jumda))+"\nK = "+str(k))
					# menentukan interval
					i = round(j / k)
					self["interval"].update(value="I = "+str(j)+" / "+str(k)+"\nI = "+str(i))
					# memasukkan self.nilai interval dan self.frekuensi
					for w in range(0, k):  # < -- Interval kelas
						if w == 0: self.kelin.append(str(nilren)+" - "+str(nilren+i-1)); psm = nilren+i-1
						else: psm += 1; self.kelin.append(str(psm)+" - "+str(psm+i-1)); psm += i-1
					psm = None
					for w in range(0, k): # < -- self.frekuensi
						if w==0:
							for u in self.nilur:
								if u in range(nilren, nilren+i): gem += 1
							self.fre.append(str(gem)); pem = nilren+i-1; gem = 0
						else:
							pem += 1
							for u in self.nilur:
								if u in range(pem, pem+i): gem += 1
							self.fre.append(str(gem)); pem += i-1; gem = 0
					for i, u in zip(self.kelin, self.fre): self.nilai.append((i, u))
					for i in self.fre: self.tofre += int(i)
					self.nilai.append(("Total", str(self.tofre)));
					self["tabel"].update(values=self.nilai)
				else: sg.popup("Inisialisasi nya diisi bro..", icon="berkas_mdc/gambar/ti.png", title="Terjadi masalah :(")
			if event in ("_s"): # <-- Simpan ke docx
				if values["_r"] != "" and values["_d"] != "" and values["_t"] != "" and values["_da"] != "":
					dok = Document()
					dok.add_heading("Distribusi frekuensi", 0)
					prg = []
					prg.append(dok.add_paragraph("Jumlah data = "+values["_d"])) # 0
					prg.append(dok.add_paragraph("self.nilai terkecil = "+values["_r"])) # 1
					prg.append(dok.add_paragraph("self.nilai terbesar = "+values["_t"])) # 2
					prg.append(dok.add_paragraph("Data acak = ")) # 3
					for i, d in enumerate(self.nila):
						if i == len(self.nila)-1: prg[3].add_run(str(d))
						else: prg[3].add_run(str(d)+" - ")
					prg.append(dok.add_paragraph("Data urut = ")) # 4
					for i, d in enumerate(self.nilur):
						if i == len(self.nilur)-1: prg[4].add_run(str(d))
						else: prg[4].add_run(str(d)+" - ")
					prg.append(dok.add_paragraph(values["jarak"])) # 5
					prg.append(dok.add_paragraph(values["kelas"])) # 5
					prg.append(dok.add_paragraph(values["interval"])) # 5
					# tabel
					t = dok.add_table(rows=1, cols=2)
					tc = t.rows[0].cells
					tc[0].text = "Interval Kelas"; tc[1].text = "frekuensi"
					for k, f in zip(self.kelin, self.fre):
						tc = t.add_row().cells
						tc[0].text = k
						tc[1].text = str(f)
					tc = t.add_row().cells; tc[0].text = "Total"; tc[1].text = str(self.tofre)
					# simpan
					ide = 0
					while True:
						if exists("berkas_mdc/ekspor/distribusi_frekuensi"+str(ide)+".docx"): ide += 1
						else: dok.save("berkas_mdc/ekspor/distribusi_frekuensi"+str(ide)+".docx");  break
					sg.popup("Berhasil menyimpan :D", icon="berkas_mdc/gambar/ti.png", title="Info")
				else: sg.popup("Data belum ada", icon="berkas_mdc/gambar/ti.png", title="Terjadi masalah :(")
			if event in ("Bersihkan"): # membersihkan data
				self["_r"].update(value=""); self["_t"].update(value=""); self["_d"].update(value="")
				self["nte"].update(value="self.nilai terendah : 0"); self["nti"].update(value="self.nilai tertinggi : 0")
				self["jd"].update(value="Jumlah data : 0"); self["_ur"].update(value=""); self["_da"].update(value="")
				self["jarak"].update(value="");self["kelas"].update(value=""); self["interval"].update(value="")
				self["tabel"].update(values="")
		
